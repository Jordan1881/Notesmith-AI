from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from pypdf import PdfReader
from docx import Document

from shared.types import ExtractResult, ExtractMeta


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> Tuple[str, int]:
    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts), len(reader.pages)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: List[str] = []

    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_document_text(file_path: str) -> ExtractResult:
    path = Path(file_path).expanduser()

    if not path.exists():
        return ExtractResult(
            text="",
            meta=ExtractMeta(
                char_count=0,
                page_estimate=0,
                extraction_success=False,
                reason=f"file_not_found: {path}"
            )
        )

    if path.is_dir():
        return ExtractResult(
            text="",
            meta=ExtractMeta(
                char_count=0,
                page_estimate=0,
                extraction_success=False,
                reason=f"expected_file_got_directory: {path}"
            )
        )

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return ExtractResult(
            text="",
            meta=ExtractMeta(
                char_count=0,
                page_estimate=0,
                extraction_success=False,
                reason=f"unsupported_extension: {ext}"
            )
        )

    try:
        if ext == ".txt":
            text = _read_txt(path)
            page_estimate = max(1, len(text) // 2500) if text.strip() else 0

        elif ext == ".pdf":
            text, page_estimate = _read_pdf(path)

        else:
            text = _read_docx(path)
            page_estimate = max(1, len(text) // 2500) if text.strip() else 0

        cleaned = _normalize_text(text)
        char_count = len(cleaned)

        if not cleaned.strip():
            reason = "empty_text_extracted (possibly scanned PDF -> needs OCR)" if ext == ".pdf" else "empty_file_or_no_text_extracted"
            return ExtractResult(
                text="",
                meta=ExtractMeta(
                    char_count=0,
                    page_estimate=page_estimate,
                    extraction_success=False,
                    reason=reason
                )
            )

        return ExtractResult(
            text=cleaned,
            meta=ExtractMeta(
                char_count=char_count,
                page_estimate=page_estimate,
                extraction_success=True,
                reason=None
            )
        )

    except Exception as e:
        return ExtractResult(
            text="",
            meta=ExtractMeta(
                char_count=0,
                page_estimate=0,
                extraction_success=False,
                reason=f"exception: {type(e).__name__}: {e}"
            )
        )


def _normalize_text(text: str) -> str:
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    while "\n\n\n" in t:
        t = t.replace("\n\n\n", "\n\n")
    return t.strip()